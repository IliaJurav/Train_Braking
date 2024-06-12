"""
Модуль экспорта результатов расчётов в doc-файл

    Выпускная квалификационная работа бакалавра
    Санкт-Петербургский политехнический университет Петра Великого
    Институт компьютерных наук и кибербезопасности
    Высшая школа компьютерных технологий и информационных систем
    направление 09.03.01 Информатика и вычислительная техника
    2024 год

@author: Журавский Илья Александрович

"""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import io
from tkinter import filedialog

# вывод таблицы в файл
def put_table(doc, df):
    t = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1], style='Table Grid')
    t.allow_autofit = True
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(df.shape[0] + 1):
        for j in range(df.shape[1]):
            if i==0:
                c = t.cell(0,j)
                p = c.paragraphs[0]
                g = df.columns[j]
                p.add_run(str(g)).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                c = t.cell(i,j)
                g = df.iat[i-1, j]
                p = c.paragraphs[0]
                p.add_run(str(g))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# ---------------------------------------------------------------------------
#
#
# name - заголовок документа
# tabls - список таблиц для экспорта
# figs - список графиков для экспорта
def putDocFile(name,tabls,figs):
    filepath = filedialog.asksaveasfilename(
        defaultextension=".docx",initialfile = "rezult.docx",title = "Save to",
        filetypes=(("docx file (.docx)", "*.docx"),("All Files", "*.*") ))
    if filepath == "":
        return
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20.0)
    section.right_margin = Mm(10.0)
    section.top_margin = Mm(15.0)
    section.bottom_margin = Mm(15.0)
    section.header_distance = Mm(10.0)
    section.footer_distance = Mm(10.0)
    def tit_head(a):
        d = document.add_paragraph('')
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = d.add_run(a)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(16)
        r.font.bold = True
# Вывод заголовка документа
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    d = document.add_paragraph('')
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run(name)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.underline = True
    d = document.add_paragraph('')
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run('Исходные данные')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.font.bold = True
    d = document.add_paragraph('')
# Вывод таблиц
    # Заголовок таблицы (брать из tabl['name'])
    # Комментарии по таблице ?
    # Собственно таблица (заголовок - tabl['cols'], данные - tabl['rows'])
    for tabl in tabls:
        tit_head(tabl['name'])
        df = pd.DataFrame(tabl['rows'], columns=tabl['cols'])
        put_table(document, df)
        d = document.add_paragraph('')

# Вывод картинки
    memfile = io.BytesIO()
    plt_rezult = figs[0]
    plt_rezult.savefig(memfile)
    memfile.seek(0)
    document.add_picture(memfile, width = Mm(180))#, width=Inches(1.25))

    document.save(filepath)
    memfile.close()
# ---------------------------------------------------------------------------